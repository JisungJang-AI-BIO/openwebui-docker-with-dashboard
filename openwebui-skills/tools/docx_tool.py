"""
title: DOCX Document Tool
author: Internal Team
author_url: https://github.com/anthropics/skills
description: Read, create, edit, and convert Word documents (.docx).
    Use when user uploads a Word document or requests new document creation.
    Supports: reading content, creating new documents with python-docx or docx-js,
    editing existing documents at XML level (unpack/edit/pack), format conversion,
    tracked changes acceptance, comments, and document validation.
required_open_webui_version: 0.8.5
requirements: python-docx, lxml, mammoth, defusedxml
version: 2.0.0
licence: MIT
"""

import asyncio
import base64
import io
import json
import os
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import Any, Callable, Optional

from pydantic import BaseModel, Field


# =============================================================================
# EventEmitter Helper
# =============================================================================
class EventEmitter:
    def __init__(self, event_emitter: Callable[[dict], Any] = None, user: dict = None):
        self.event_emitter = event_emitter
        self.user = user or {}

    async def emit(self, description="Unknown State", status="in_progress", done=False):
        if self.event_emitter:
            await self.event_emitter(
                {"type": "status", "data": {"status": status, "description": description, "done": done}}
            )

    async def progress(self, description: str):
        await self.emit(description)

    async def success(self, description: str):
        await self.emit(description, "success", True)

    async def error(self, description: str):
        await self.emit(description, "error", True)

    async def send_file_link(self, file_path: str, filename: str, mime_type: str = None):
        """Upload file to OpenWebUI Files API and emit download link.
        Falls back to base64 data URI if the API is unavailable.
        """
        if not self.event_emitter or not os.path.exists(file_path):
            return

        if mime_type is None:
            ext = Path(filename).suffix.lower()
            mime_map = {
                ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
                ".pdf": "application/pdf", ".html": "text/html", ".txt": "text/plain",
                ".csv": "text/csv", ".odt": "application/vnd.oasis.opendocument.text",
                ".rtf": "application/rtf", ".png": "image/png",
                ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".gif": "image/gif",
            }
            mime_type = mime_map.get(ext, "application/octet-stream")

        is_image = mime_type.startswith("image/")

        # Strategy 1: Upload to OpenWebUI Files API
        url = await self._upload_to_openwebui(file_path, filename, mime_type)

        if url:
            if is_image:
                content = f"\n\n📎 **{filename}**\n\n![{filename}]({url})\n\n[Download {filename}]({url})\n"
            else:
                content = f"\n\n📎 **{filename}**\n\n[Download {filename}]({url})\n"
        else:
            # Fallback: base64 data URI
            with open(file_path, "rb") as f:
                b64 = base64.b64encode(f.read()).decode("utf-8")
            data_uri = f"data:{mime_type};base64,{b64}"
            if is_image:
                content = f"\n\n📎 **{filename}**\n\n![{filename}]({data_uri})\n\n[Download {filename}]({data_uri})\n"
            else:
                content = f"\n\n📎 **{filename}**\n\n[Download {filename}]({data_uri})\n"

        await self.event_emitter({"type": "message", "data": {"content": content}})

    async def _upload_to_openwebui(self, file_path: str, filename: str, mime_type: str) -> Optional[str]:
        """Upload file via OpenWebUI's internal Files API. Returns download URL or None."""
        try:
            import httpx
            import jwt

            user_id = self.user.get("id", "")
            secret = os.environ.get("WEBUI_SECRET_KEY", "")
            if not user_id or not secret:
                return None

            token = jwt.encode({"id": user_id}, secret, algorithm="HS256")
            if isinstance(token, bytes):
                token = token.decode("utf-8")

            port = os.environ.get("PORT", "8080")
            with open(file_path, "rb") as f:
                resp = httpx.post(
                    f"http://localhost:{port}/api/v1/files/",
                    files={"file": (filename, f, mime_type)},
                    headers={"Authorization": f"Bearer {token}"},
                    timeout=30.0,
                )

            if resp.status_code in (200, 201):
                file_id = resp.json().get("id")
                if file_id:
                    return f"/api/v1/files/{file_id}/content"
            return None
        except Exception:
            return None


# =============================================================================
# Main Tool Class
# =============================================================================
class Tools:
    class Valves(BaseModel):
        TEMP_DIR: str = Field(
            default="/tmp/openwebui-docx",
            description="Temporary file storage path for document processing",
        )
        MAX_FILE_SIZE_MB: int = Field(
            default=50, description="Maximum file size in MB"
        )
        SCRIPTS_DIR: str = Field(
            default="",
            description="Absolute path to Anthropic docx scripts (e.g. /data/OpenWebUI-Skills/vendor/docx). Contains office/, accept_changes.py, comment.py. Leave empty for fallback mode.",
        )
        LIBREOFFICE_PATH: str = Field(
            default="soffice",
            description="LibreOffice binary path",
        )
        PANDOC_PATH: str = Field(
            default="pandoc",
            description="Pandoc binary path",
        )
        NODE_PATH: str = Field(
            default="node",
            description="Node.js binary path (for docx-js document creation)",
        )
        DOCXJS_REQUIRE_PATH: str = Field(
            default="docx",
            description="docx-js require path (global 'docx' or absolute path to node_modules)",
        )
        USE_DOCXJS: bool = Field(
            default=True,
            description="Use docx-js (Node.js) for document creation. If False, falls back to python-docx.",
        )
        DEFAULT_FONT: str = Field(
            default="Arial",
            description="Default font for new documents (python-docx fallback)",
        )

    class UserValves(BaseModel):
        AUTHOR_NAME: str = Field(
            default="User",
            description="Author name for document metadata and tracked changes",
        )

    def __init__(self):
        self.valves = self.Valves()
        self.citation = False

    # =========================================================================
    # Tool Methods (exposed to LLM via function calling)
    # =========================================================================

    async def read_docx(
        self,
        mode: str = "text",
        __files__: list = None,
        __user__: dict = None,
        __event_emitter__: Callable[[dict], Any] = None,
    ) -> str:
        """
        Read and extract content from an uploaded Word document (.docx).
        Supported modes:
        - "text": Plain text extraction with headings, lists, and tables
        - "structured": Detailed structure (sections, styles, paragraph indices)
        - "xml": Raw document.xml content (for debugging/XML editing)
        - "markdown": Pandoc-based Markdown conversion (if pandoc available)

        :param mode: Extraction mode - "text", "structured", "xml", or "markdown"
        :return: Extracted document content
        """
        emitter = EventEmitter(__event_emitter__, __user__)
        await emitter.progress("Reading DOCX file...")

        if not __files__:
            await emitter.error("No file uploaded")
            return "Error: No file uploaded. Please upload a .docx file."

        file_path = self._resolve_docx_file(__files__)
        if not file_path:
            await emitter.error("No .docx file found")
            return "Error: No .docx file found in uploaded files."

        try:
            if mode == "xml":
                result = self._read_xml(file_path)
            elif mode == "structured":
                result = self._read_structured(file_path)
            elif mode == "markdown":
                result = await self._read_with_pandoc(file_path, emitter)
            else:
                result = self._read_text(file_path)

            await emitter.success(f"Read {Path(file_path).name}")
            return result

        except Exception as e:
            await emitter.error(f"Read error: {str(e)}")
            return f"Error reading DOCX: {str(e)}"

    async def create_docx(
        self,
        content_description: str,
        filename: str = "document.docx",
        page_size: str = "A4",
        orientation: str = "portrait",
        __user__: dict = None,
        __event_emitter__: Callable[[dict], Any] = None,
    ) -> str:
        """
        Create a new Word document. There are two modes:

        MODE 1 (docx-js enabled / recommended):
        Provide a content_description of what the document should contain.
        You (the LLM) should ALSO generate JavaScript code using the docx library
        and pass it to create_docx_from_js() for complex documents with precise formatting.

        MODE 2 (python-docx fallback):
        Provide content in Markdown-like format:
        - '# ' → Heading 1, '## ' → Heading 2, '### ' → Heading 3
        - '- ' → bullet list, '1. ' → numbered list
        - '| col1 | col2 |' → table rows (first row = header, skip separator rows)
        - '---' → page break
        - '> ' → block quote
        - Other lines → normal paragraphs

        :param content_description: Document content specification
        :param filename: Output filename (default: document.docx)
        :param page_size: "A4", "Letter", or "Legal"
        :param orientation: "portrait" or "landscape"
        :return: Status message with download link
        """
        emitter = EventEmitter(__event_emitter__, __user__)
        await emitter.progress("Creating DOCX...")

        try:
            temp_dir = self._ensure_temp_dir()
            output_path = os.path.join(temp_dir, filename)

            # python-docx fallback
            result = self._create_with_python_docx(
                content_description, output_path, page_size, orientation, __user__
            )

            if "Error" in result:
                await emitter.error(result)
                return result

            await emitter.send_file_link(output_path, filename)
            await emitter.success(f"Created {filename}")

            file_size = os.path.getsize(output_path)
            return (
                f"Created '{filename}' ({file_size:,} bytes). "
                f"Page: {page_size} {orientation}. Download link above."
            )

        except Exception as e:
            await emitter.error(f"Creation error: {str(e)}")
            return f"Error creating DOCX: {str(e)}"

    async def create_docx_from_js(
        self,
        js_code: str,
        filename: str = "document.docx",
        __user__: dict = None,
        __event_emitter__: Callable[[dict], Any] = None,
    ) -> str:
        """
        Create a Word document by executing docx-js (JavaScript) code via Node.js.
        This method gives full control over document layout using the docx npm library.

        The js_code should:
        1. Use `const { Document, Packer, ... } = require('docx')` (auto-injected)
        2. Create a `doc` variable as a new Document
        3. NOT include the Packer/write step (auto-injected)

        Example js_code:
        ```
        const doc = new Document({
            sections: [{
                properties: {
                    page: { size: { width: 12240, height: 15840 } }
                },
                children: [
                    new Paragraph({
                        heading: HeadingLevel.HEADING_1,
                        children: [new TextRun("Hello World")]
                    })
                ]
            }]
        });
        ```

        :param js_code: JavaScript code that creates a `doc` variable using docx library
        :param filename: Output filename (default: document.docx)
        :return: Status message with download link
        """
        emitter = EventEmitter(__event_emitter__, __user__)

        if not self.valves.USE_DOCXJS:
            await emitter.error("docx-js disabled in Valves")
            return "Error: docx-js is disabled. Set USE_DOCXJS=True in Valves or use create_docx() instead."

        await emitter.progress("Executing docx-js via Node.js...")

        try:
            temp_dir = self._ensure_temp_dir()
            output_path = os.path.join(temp_dir, filename)
            js_file = os.path.join(temp_dir, "_create_doc.js")

            # Build full JS with require and output
            output_path_escaped = output_path.replace("\\", "/")
            require_path = self.valves.DOCXJS_REQUIRE_PATH

            full_js = f"""
const {{ Document, Packer, Paragraph, TextRun, Table, TableRow, TableCell,
  ImageRun, Header, Footer, AlignmentType, PageOrientation, LevelFormat,
  ExternalHyperlink, TableOfContents, HeadingLevel, BorderStyle, WidthType,
  ShadingType, VerticalAlign, PageNumber, PageBreak, Tab, TabStopType,
  TabStopPosition, NumberFormat, convertInchesToTwip }} = require('{require_path}');
const fs = require('fs');

// --- User code starts ---
{js_code}
// --- User code ends ---

// Write output
Packer.toBuffer(doc).then(buffer => {{
    fs.writeFileSync("{output_path_escaped}", buffer);
    console.log(JSON.stringify({{success: true, size: buffer.length}}));
}}).catch(err => {{
    console.error(JSON.stringify({{success: false, error: err.message}}));
    process.exit(1);
}});
"""
            with open(js_file, "w", encoding="utf-8") as f:
                f.write(full_js)

            result = subprocess.run(
                [self.valves.NODE_PATH, js_file],
                capture_output=True,
                text=True,
                timeout=30,
                cwd=temp_dir,
            )

            if result.returncode != 0:
                stderr = result.stderr.strip()
                await emitter.error("docx-js execution failed")
                return f"Error: docx-js execution failed.\nStderr: {stderr}\nStdout: {result.stdout.strip()}"

            if not os.path.exists(output_path):
                await emitter.error("Output file not created")
                return f"Error: docx-js ran but output file was not created.\nStdout: {result.stdout}\nStderr: {result.stderr}"

            # Validate
            await emitter.progress("Validating generated document...")
            validation_errors = self._validate_structure(output_path)
            validation_msg = ""
            if validation_errors:
                validation_msg = f"\n⚠️ Validation warnings: {'; '.join(validation_errors[:3])}"

            await emitter.send_file_link(output_path, filename)
            await emitter.success(f"Created {filename}")

            file_size = os.path.getsize(output_path)
            return (
                f"Created '{filename}' ({file_size:,} bytes) via docx-js.{validation_msg}\n"
                f"Download link above."
            )

        except subprocess.TimeoutExpired:
            await emitter.error("Node.js execution timed out")
            return "Error: docx-js execution timed out after 30 seconds."
        except FileNotFoundError:
            await emitter.error("Node.js not found")
            return (
                f"Error: Node.js not found at '{self.valves.NODE_PATH}'. "
                f"Install Node.js or update NODE_PATH in Valves."
            )
        except Exception as e:
            await emitter.error(f"Creation error: {str(e)}")
            return f"Error: {str(e)}"

    async def edit_docx_xml(
        self,
        operation: str,
        parameters: str = "",
        __files__: list = None,
        __user__: dict = None,
        __event_emitter__: Callable[[dict], Any] = None,
    ) -> str:
        """
        Edit an uploaded DOCX at the XML level using unpack → edit → pack workflow.
        This is the most powerful editing method, operating directly on OOXML.

        Operations:
        - "unpack": Unpack to inspect XML structure. Returns document.xml content.
        - "replace_text": Replace text. parameters = "old_text|||new_text"
        - "replace_xml": Replace XML fragment. parameters = "old_xml|||new_xml"
          (operates on word/document.xml)
        - "accept_changes": Accept all tracked changes via LibreOffice.
        - "add_comment": Add a comment. parameters = "comment_id|||text|||xpath_hint"
          (xpath_hint is optional, describes where to place the comment)
        - "repack": After manual XML edits, repack into .docx.
          parameters = unpacked directory path (from previous unpack)

        :param operation: Operation name
        :param parameters: Operation-specific parameters
        :return: Result with XML content or download link
        """
        emitter = EventEmitter(__event_emitter__, __user__)
        await emitter.progress(f"DOCX XML edit: {operation}...")

        if operation not in ("repack",) and not __files__:
            await emitter.error("No file uploaded")
            return "Error: No file uploaded."

        try:
            file_path = self._resolve_docx_file(__files__) if __files__ else None
            temp_dir = self._ensure_temp_dir()
            scripts_dir = self.valves.SCRIPTS_DIR

            if operation == "unpack":
                return await self._op_unpack(file_path, temp_dir, scripts_dir, emitter)
            elif operation == "replace_text":
                return await self._op_replace_text(file_path, parameters, temp_dir, scripts_dir, emitter)
            elif operation == "replace_xml":
                return await self._op_replace_xml(file_path, parameters, temp_dir, scripts_dir, emitter)
            elif operation == "accept_changes":
                return await self._op_accept_changes(file_path, temp_dir, scripts_dir, emitter)
            elif operation == "add_comment":
                return await self._op_add_comment(file_path, parameters, temp_dir, scripts_dir, emitter)
            elif operation == "repack":
                return await self._op_repack(parameters, temp_dir, scripts_dir, emitter)
            else:
                return f"Error: Unknown operation '{operation}'. Available: unpack, replace_text, replace_xml, accept_changes, add_comment, repack"

        except Exception as e:
            await emitter.error(f"Edit error: {str(e)}")
            return f"Error: {str(e)}"

    async def convert_docx(
        self,
        target_format: str = "pdf",
        __files__: list = None,
        __user__: dict = None,
        __event_emitter__: Callable[[dict], Any] = None,
    ) -> str:
        """
        Convert an uploaded Word document to another format.
        Uses LibreOffice (pdf/odt/rtf) or Pandoc (html/txt/markdown).

        :param target_format: "pdf", "html", "txt", "odt", "rtf", "md" (markdown)
        :return: Converted file download or text content
        """
        emitter = EventEmitter(__event_emitter__, __user__)
        await emitter.progress(f"Converting to {target_format}...")

        if not __files__:
            await emitter.error("No file uploaded")
            return "Error: No file uploaded."

        file_path = self._resolve_docx_file(__files__)
        if not file_path:
            await emitter.error("No .docx file found")
            return "Error: No .docx file found."

        try:
            temp_dir = self._ensure_temp_dir()
            base_name = Path(file_path).stem
            fmt = target_format.lower().strip(".")

            if fmt in ("pdf", "odt", "rtf"):
                return await self._convert_libreoffice(file_path, temp_dir, fmt, emitter)
            elif fmt in ("html", "txt", "md", "plain", "markdown"):
                pandoc_fmt = {"txt": "plain", "md": "markdown"}.get(fmt, fmt)
                return await self._convert_pandoc(file_path, temp_dir, pandoc_fmt, emitter)
            else:
                return f"Error: Unsupported format '{fmt}'. Supported: pdf, html, txt, odt, rtf, md"

        except Exception as e:
            await emitter.error(f"Conversion error: {str(e)}")
            return f"Error: {str(e)}"

    async def validate_docx(
        self,
        __files__: list = None,
        __user__: dict = None,
        __event_emitter__: Callable[[dict], Any] = None,
    ) -> str:
        """
        Validate DOCX structure: XML well-formedness, relationships, content types.
        If the Anthropic scripts are available, uses the full XSD schema validation.

        :return: Validation report
        """
        emitter = EventEmitter(__event_emitter__, __user__)
        await emitter.progress("Validating DOCX...")

        if not __files__:
            await emitter.error("No file uploaded")
            return "Error: No file uploaded."

        file_path = self._resolve_docx_file(__files__)
        if not file_path:
            await emitter.error("No .docx file found")
            return "Error: No .docx file found."

        try:
            scripts_dir = self.valves.SCRIPTS_DIR
            validate_script = os.path.join(scripts_dir, "office", "validate.py")

            if os.path.exists(validate_script):
                # Use Anthropic's full validation
                result = subprocess.run(
                    ["python", validate_script, file_path, "--auto-repair"],
                    capture_output=True,
                    text=True,
                    timeout=30,
                    cwd=os.path.join(scripts_dir, "office"),
                )
                output = (result.stdout + "\n" + result.stderr).strip()
                if result.returncode == 0:
                    await emitter.success("Validation passed")
                    return f"✅ Validation PASSED\n\n{output}"
                else:
                    await emitter.error("Validation failed")
                    return f"❌ Validation FAILED\n\n{output}"
            else:
                # Fallback: basic structural validation
                errors = self._validate_structure(file_path)
                if not errors:
                    await emitter.success("Basic validation passed")
                    return "✅ Basic validation PASSED (XML well-formed, required files present).\nNote: Full XSD schema validation requires Anthropic scripts at SCRIPTS_DIR."
                else:
                    await emitter.error(f"{len(errors)} error(s) found")
                    return "❌ Validation FAILED:\n" + "\n".join(f"- {e}" for e in errors)

        except Exception as e:
            await emitter.error(f"Validation error: {str(e)}")
            return f"Error: {str(e)}"

    # =========================================================================
    # Internal: File Resolution
    # =========================================================================

    def _resolve_docx_file(self, files: list) -> Optional[str]:
        """Find and resolve path for first .docx file in uploaded files."""
        if not files:
            return None

        for f in files:
            info = f if isinstance(f, dict) else {"path": f}
            name = info.get("filename", info.get("name", str(info.get("path", ""))))

            if not name.lower().endswith((".docx", ".doc")):
                continue

            for key in ("path", "file_path", "url", "id"):
                path = info.get(key)
                if path and isinstance(path, str):
                    candidates = [
                        path,
                        f"/app/backend/data/uploads/{path}",
                        f"/app/backend/data/cache/{path}",
                    ]
                    for c in candidates:
                        if os.path.exists(c):
                            return c
                    return path  # return as-is, let caller handle missing

        # Fallback: return first file regardless of extension
        if files:
            f = files[0]
            if isinstance(f, dict):
                for key in ("path", "file_path"):
                    if key in f:
                        return f[key]
            elif isinstance(f, str):
                return f
        return None

    def _ensure_temp_dir(self) -> str:
        os.makedirs(self.valves.TEMP_DIR, exist_ok=True)
        return self.valves.TEMP_DIR

    # =========================================================================
    # Internal: Reading
    # =========================================================================

    def _read_text(self, file_path: str) -> str:
        from docx import Document

        doc = Document(file_path)
        parts = []

        props = doc.core_properties
        if props.title:
            parts.append(f"**Title:** {props.title}")
        if props.author:
            parts.append(f"**Author:** {props.author}")
        if parts:
            parts.append("")

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                parts.append("")
                continue

            style_name = para.style.name if para.style else "Normal"
            if "Heading" in style_name:
                try:
                    level = int(style_name.replace("Heading ", ""))
                    parts.append(f"{'#' * level} {text}")
                except ValueError:
                    parts.append(f"**{text}**")
            elif "List" in style_name or "Bullet" in style_name:
                parts.append(f"- {text}")
            else:
                parts.append(text)

        for i, table in enumerate(doc.tables):
            parts.append(f"\n**Table {i+1}:**")
            for j, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                parts.append("| " + " | ".join(cells) + " |")
                if j == 0:
                    parts.append("| " + " | ".join(["---"] * len(cells)) + " |")

        return "\n".join(parts)

    def _read_structured(self, file_path: str) -> str:
        from docx import Document

        doc = Document(file_path)
        parts = ["## Document Structure\n"]

        parts.append(f"**Sections:** {len(doc.sections)}")
        for i, section in enumerate(doc.sections):
            w, h = section.page_width, section.page_height
            parts.append(f"  Section {i+1}: {w/914400:.1f}\" × {h/914400:.1f}\"")

        parts.append(f"**Paragraphs:** {len(doc.paragraphs)}")
        parts.append(f"**Tables:** {len(doc.tables)}")

        styles_used = {p.style.name for p in doc.paragraphs if p.style}
        parts.append(f"**Styles:** {', '.join(sorted(styles_used))}")

        parts.append("\n## Content\n")
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            style = para.style.name if para.style else "Normal"
            if text:
                parts.append(f"[{i}] ({style}) {text}")

        for i, table in enumerate(doc.tables):
            parts.append(f"\n### Table {i+1} ({len(table.rows)}×{len(table.columns)})")
            for j, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(f"  Row {j}: | " + " | ".join(cells) + " |")

        return "\n".join(parts)

    def _read_xml(self, file_path: str) -> str:
        import defusedxml.minidom

        with zipfile.ZipFile(file_path, "r") as zf:
            if "word/document.xml" not in zf.namelist():
                return "Error: word/document.xml not found in archive."
            xml_bytes = zf.read("word/document.xml")
            dom = defusedxml.minidom.parseString(xml_bytes)
            pretty = dom.toprettyxml(indent="  ")
            if len(pretty) > 20000:
                return pretty[:20000] + "\n\n... (truncated at 20000 chars)"
            return pretty

    async def _read_with_pandoc(self, file_path: str, emitter: EventEmitter) -> str:
        try:
            result = subprocess.run(
                [self.valves.PANDOC_PATH, "--track-changes=all", file_path, "-t", "markdown"],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0:
                return result.stdout
            return f"Pandoc error: {result.stderr}"
        except FileNotFoundError:
            return f"Error: Pandoc not found at '{self.valves.PANDOC_PATH}'."

    # =========================================================================
    # Internal: Creating (python-docx fallback)
    # =========================================================================

    def _create_with_python_docx(
        self, content_spec: str, output_path: str, page_size: str, orientation: str, user: dict
    ) -> str:
        from docx import Document
        from docx.shared import Inches, Mm, Pt
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()
        section = doc.sections[0]

        # Page setup
        sizes = {"A4": (Mm(210), Mm(297)), "Letter": (Inches(8.5), Inches(11)), "Legal": (Inches(8.5), Inches(14))}
        width, height = sizes.get(page_size, sizes["A4"])
        if orientation.lower() == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = height, width
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width, section.page_height = width, height
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)

        # Default style
        style = doc.styles["Normal"]
        style.font.name = self.valves.DEFAULT_FONT
        style.font.size = Pt(12)

        # Build content
        self._build_content(doc, content_spec)

        # Metadata
        author = "User"
        if user:
            author = user.get("name", user.get("valves", {}).get("AUTHOR_NAME", "User")) if isinstance(user, dict) else "User"
        doc.core_properties.author = author

        doc.save(output_path)
        return "OK"

    def _build_content(self, doc, spec: str):
        from docx import Document
        from docx.shared import Inches

        table_buf = []

        for line in spec.split("\n"):
            s = line.strip()

            # Table accumulation
            if s.startswith("|") and s.endswith("|"):
                inner = s[1:-1].strip()
                if all(c in "-| :" for c in inner):
                    continue  # separator row
                table_buf.append(s)
                continue
            elif table_buf:
                self._flush_table(doc, table_buf)
                table_buf = []

            if s == "---":
                doc.add_page_break()
            elif not s:
                doc.add_paragraph("")
            elif s.startswith("### "):
                doc.add_heading(s[4:], level=3)
            elif s.startswith("## "):
                doc.add_heading(s[3:], level=2)
            elif s.startswith("# "):
                doc.add_heading(s[2:], level=1)
            elif s.startswith("- ") or s.startswith("* "):
                doc.add_paragraph(s[2:], style="List Bullet")
            elif len(s) > 2 and s[0].isdigit() and (s[1] == "." or (s[1].isdigit() and s[2] == ".")):
                text = s.split(".", 1)[1].strip()
                doc.add_paragraph(text, style="List Number")
            elif s.startswith("> "):
                p = doc.add_paragraph(s[2:])
                p.paragraph_format.left_indent = Inches(0.5)
                for r in p.runs:
                    r.italic = True
            elif s.startswith("**") and s.endswith("**"):
                p = doc.add_paragraph()
                p.add_run(s[2:-2]).bold = True
            else:
                doc.add_paragraph(s)

        if table_buf:
            self._flush_table(doc, table_buf)

    def _flush_table(self, doc, buf: list):
        from docx.enum.table import WD_TABLE_ALIGNMENT

        rows_data = [[c.strip() for c in row.strip("|").split("|")] for row in buf]
        if not rows_data:
            return
        ncols = max(len(r) for r in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=ncols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, rd in enumerate(rows_data):
            for j, ct in enumerate(rd):
                if j < ncols:
                    table.cell(i, j).text = ct
                    if i == 0:
                        for p in table.cell(i, j).paragraphs:
                            for r in p.runs:
                                r.bold = True
        doc.add_paragraph("")

    # =========================================================================
    # Internal: XML Edit Operations (using Anthropic scripts where possible)
    # =========================================================================

    async def _op_unpack(self, file_path, temp_dir, scripts_dir, emitter) -> str:
        unpack_dir = os.path.join(temp_dir, "unpacked")
        unpack_script = os.path.join(scripts_dir, "office", "unpack.py")

        if os.path.exists(unpack_script):
            if os.path.exists(unpack_dir):
                shutil.rmtree(unpack_dir)
            result = subprocess.run(
                ["python", unpack_script, file_path, unpack_dir],
                capture_output=True, text=True, timeout=30,
                cwd=os.path.join(scripts_dir, "office"),
            )
            output = result.stdout.strip()
        else:
            # Fallback: manual unpack
            if os.path.exists(unpack_dir):
                shutil.rmtree(unpack_dir)
            os.makedirs(unpack_dir)
            import defusedxml.minidom
            with zipfile.ZipFile(file_path, "r") as zf:
                zf.extractall(unpack_dir)
            for xf in Path(unpack_dir).rglob("*.xml"):
                try:
                    dom = defusedxml.minidom.parseString(xf.read_bytes())
                    xf.write_bytes(dom.toprettyxml(indent="  ", encoding="utf-8"))
                except Exception:
                    pass
            output = "Unpacked (fallback mode, no run merging/redline simplification)"

        # Read document.xml
        doc_xml = os.path.join(unpack_dir, "word", "document.xml")
        xml_content = ""
        if os.path.exists(doc_xml):
            xml_content = Path(doc_xml).read_text(encoding="utf-8")

        file_list = [str(p.relative_to(unpack_dir)) for p in Path(unpack_dir).rglob("*") if p.is_file()]

        await emitter.success("Unpacked")
        result_text = f"{output}\n\n**Unpacked to:** `{unpack_dir}`\n\n"
        result_text += f"**Files ({len(file_list)}):**\n```\n" + "\n".join(sorted(file_list)) + "\n```\n\n"
        result_text += f"**document.xml:**\n```xml\n{xml_content[:15000]}\n```"
        if len(xml_content) > 15000:
            result_text += "\n(truncated)"
        return result_text

    async def _op_replace_text(self, file_path, parameters, temp_dir, scripts_dir, emitter) -> str:
        parts = parameters.split("|||")
        if len(parts) != 2:
            return "Error: format is 'old_text|||new_text'"
        old_text, new_text = parts

        # Unpack → replace in XML → repack
        unpack_dir = os.path.join(temp_dir, "edit_unpack")
        self._simple_unpack(file_path, unpack_dir)

        doc_xml = os.path.join(unpack_dir, "word", "document.xml")
        if not os.path.exists(doc_xml):
            return "Error: document.xml not found"

        content = Path(doc_xml).read_text(encoding="utf-8")
        count = content.count(old_text)
        if count == 0:
            # Try across XML tags (text split across runs)
            from docx import Document
            doc = Document(file_path)
            base_name = Path(file_path).stem
            output_path = os.path.join(temp_dir, f"{base_name}_edited.docx")
            replaced = 0
            for para in doc.paragraphs:
                if old_text in para.text:
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            replaced += 1
            doc.save(output_path)
            await emitter.send_file_link(output_path, f"{base_name}_edited.docx")
            await emitter.success("Text replaced (python-docx)")
            return f"Replaced {replaced} occurrence(s) via python-docx (text was split across XML runs)."

        content = content.replace(old_text, new_text)
        Path(doc_xml).write_text(content, encoding="utf-8")

        output_path = self._simple_repack(unpack_dir, temp_dir, Path(file_path).stem + "_edited.docx")
        await emitter.send_file_link(output_path, Path(output_path).name)
        await emitter.success("Text replaced")
        return f"Replaced {count} occurrence(s) of '{old_text}' in XML."

    async def _op_replace_xml(self, file_path, parameters, temp_dir, scripts_dir, emitter) -> str:
        parts = parameters.split("|||")
        if len(parts) != 2:
            return "Error: format is 'old_xml|||new_xml'"

        unpack_dir = os.path.join(temp_dir, "edit_unpack")
        self._simple_unpack(file_path, unpack_dir)

        doc_xml = os.path.join(unpack_dir, "word", "document.xml")
        content = Path(doc_xml).read_text(encoding="utf-8")
        count = content.count(parts[0])
        content = content.replace(parts[0], parts[1])
        Path(doc_xml).write_text(content, encoding="utf-8")

        output_path = self._simple_repack(unpack_dir, temp_dir, Path(file_path).stem + "_edited.docx")
        await emitter.send_file_link(output_path, Path(output_path).name)
        await emitter.success("XML replaced")
        return f"Replaced {count} XML fragment(s)."

    async def _op_accept_changes(self, file_path, temp_dir, scripts_dir, emitter) -> str:
        accept_script = os.path.join(scripts_dir, "accept_changes.py")
        base_name = Path(file_path).stem
        output_path = os.path.join(temp_dir, f"{base_name}_accepted.docx")

        if os.path.exists(accept_script):
            result = subprocess.run(
                ["python", accept_script, file_path, output_path],
                capture_output=True, text=True, timeout=60,
                cwd=scripts_dir,
            )
            if os.path.exists(output_path):
                await emitter.send_file_link(output_path, f"{base_name}_accepted.docx")
                await emitter.success("Changes accepted")
                return f"Accepted all tracked changes. Output: {base_name}_accepted.docx\n{result.stdout}"
            return f"Error: {result.stdout}\n{result.stderr}"
        else:
            # Fallback: direct LibreOffice call
            shutil.copy2(file_path, output_path)
            soffice = self.valves.LIBREOFFICE_PATH
            profile = os.path.join(temp_dir, "lo_profile")
            macro_dir = os.path.join(profile, "user", "basic", "Standard")
            os.makedirs(macro_dir, exist_ok=True)
            Path(os.path.join(macro_dir, "Module1.xba")).write_text(
                '<?xml version="1.0" encoding="UTF-8"?>\n'
                '<!DOCTYPE script:module PUBLIC "-//OpenOffice.org//DTD OfficeDocument 1.0//EN" "module.dtd">\n'
                '<script:module xmlns:script="http://openoffice.org/2000/script" script:name="Module1" script:language="StarBasic">\n'
                'Sub AcceptAllTrackedChanges()\n'
                '    Dim document As Object\n    Dim dispatcher As Object\n'
                '    document = ThisComponent.CurrentController.Frame\n'
                '    dispatcher = createUnoService("com.sun.star.frame.DispatchHelper")\n'
                '    dispatcher.executeDispatch(document, ".uno:AcceptAllTrackedChanges", "", 0, Array())\n'
                '    ThisComponent.store()\n    ThisComponent.close(True)\nEnd Sub\n'
                '</script:module>'
            )
            env = os.environ.copy()
            env["SAL_USE_VCLPLUGIN"] = "svp"
            try:
                subprocess.run(
                    [soffice, "--headless", f"-env:UserInstallation=file://{profile}",
                     "--norestore",
                     "vnd.sun.star.script:Standard.Module1.AcceptAllTrackedChanges?language=Basic&location=application",
                     str(Path(output_path).absolute())],
                    capture_output=True, text=True, timeout=30, env=env,
                )
            except (subprocess.TimeoutExpired, FileNotFoundError) as e:
                if isinstance(e, FileNotFoundError):
                    return f"Error: LibreOffice not found at '{soffice}'."

            if os.path.exists(output_path):
                await emitter.send_file_link(output_path, f"{base_name}_accepted.docx")
                await emitter.success("Changes accepted (fallback)")
                return f"Accepted tracked changes (fallback mode)."
            return "Error: Failed to accept changes."

    async def _op_add_comment(self, file_path, parameters, temp_dir, scripts_dir, emitter) -> str:
        comment_script = os.path.join(scripts_dir, "comment.py")
        if not os.path.exists(comment_script):
            return "Error: comment.py not found at SCRIPTS_DIR. Comment addition requires Anthropic scripts."

        parts = parameters.split("|||")
        if len(parts) < 2:
            return "Error: format is 'comment_id|||text' or 'comment_id|||text|||parent_id'"

        unpack_dir = os.path.join(temp_dir, "comment_unpack")
        self._simple_unpack(file_path, unpack_dir)

        cmd = ["python", comment_script, unpack_dir, parts[0], parts[1]]
        if len(parts) >= 3:
            cmd.extend(["--parent", parts[2]])

        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30, cwd=scripts_dir)

        output_path = self._simple_repack(unpack_dir, temp_dir, Path(file_path).stem + "_commented.docx")
        await emitter.send_file_link(output_path, Path(output_path).name)
        await emitter.success("Comment added")
        return f"Comment added.\n{result.stdout}\n\nNote: You still need to add comment markers to document.xml. See output above for marker XML."

    async def _op_repack(self, parameters, temp_dir, scripts_dir, emitter) -> str:
        unpack_dir = parameters.strip() or os.path.join(temp_dir, "unpacked")
        if not os.path.isdir(unpack_dir):
            return f"Error: Directory not found: {unpack_dir}"

        pack_script = os.path.join(scripts_dir, "office", "pack.py")
        output_path = os.path.join(temp_dir, "repacked.docx")

        if os.path.exists(pack_script):
            result = subprocess.run(
                ["python", pack_script, unpack_dir, output_path],
                capture_output=True, text=True, timeout=30,
                cwd=os.path.join(scripts_dir, "office"),
            )
            if os.path.exists(output_path):
                await emitter.send_file_link(output_path, "repacked.docx")
                await emitter.success("Repacked")
                return f"Repacked successfully.\n{result.stdout}"
            return f"Error: {result.stdout}\n{result.stderr}"
        else:
            output_path = self._simple_repack(unpack_dir, temp_dir, "repacked.docx")
            await emitter.send_file_link(output_path, "repacked.docx")
            await emitter.success("Repacked (fallback)")
            return "Repacked (fallback mode, no validation)."

    # =========================================================================
    # Internal: Conversion
    # =========================================================================

    async def _convert_libreoffice(self, file_path, temp_dir, fmt, emitter) -> str:
        soffice = self.valves.LIBREOFFICE_PATH
        input_copy = os.path.join(temp_dir, Path(file_path).name)
        shutil.copy2(file_path, input_copy)
        env = os.environ.copy()
        env["SAL_USE_VCLPLUGIN"] = "svp"
        try:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", fmt, "--outdir", temp_dir, input_copy],
                capture_output=True, text=True, timeout=60, env=env,
            )
        except FileNotFoundError:
            return f"Error: LibreOffice not found at '{soffice}'."
        except subprocess.TimeoutExpired:
            return "Error: Conversion timed out."

        output_path = os.path.join(temp_dir, f"{Path(file_path).stem}.{fmt}")
        if os.path.exists(output_path):
            await emitter.send_file_link(output_path, Path(output_path).name)
            await emitter.success(f"Converted to {fmt}")
            return f"Converted to {fmt} ({os.path.getsize(output_path):,} bytes). Download above."
        return f"Error: Conversion failed. {result.stderr}"

    async def _convert_pandoc(self, file_path, temp_dir, fmt, emitter) -> str:
        ext_map = {"plain": "txt", "markdown": "md", "html": "html"}
        ext = ext_map.get(fmt, fmt)
        output_path = os.path.join(temp_dir, f"{Path(file_path).stem}.{ext}")
        try:
            result = subprocess.run(
                [self.valves.PANDOC_PATH, "-f", "docx", "-t", fmt, "-o", output_path, file_path],
                capture_output=True, text=True, timeout=60,
            )
        except FileNotFoundError:
            return f"Error: Pandoc not found at '{self.valves.PANDOC_PATH}'."
        if os.path.exists(output_path) and result.returncode == 0:
            content = Path(output_path).read_text(encoding="utf-8", errors="replace")
            await emitter.success(f"Converted to {ext}")
            if len(content) > 10000:
                return f"## Converted to {ext}\n\n{content[:10000]}\n\n... (truncated)"
            return f"## Converted to {ext}\n\n{content}"
        return f"Error: {result.stderr}"

    # =========================================================================
    # Internal: Utility
    # =========================================================================

    def _simple_unpack(self, file_path: str, unpack_dir: str):
        if os.path.exists(unpack_dir):
            shutil.rmtree(unpack_dir)
        os.makedirs(unpack_dir)
        with zipfile.ZipFile(file_path, "r") as zf:
            zf.extractall(unpack_dir)

    def _simple_repack(self, unpack_dir: str, temp_dir: str, filename: str) -> str:
        output_path = os.path.join(temp_dir, filename)
        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for f in Path(unpack_dir).rglob("*"):
                if f.is_file():
                    zf.write(f, f.relative_to(unpack_dir))
        return output_path

    def _validate_structure(self, file_path: str) -> list:
        """Basic structural validation (no XSD)."""
        import lxml.etree
        errors = []
        try:
            with zipfile.ZipFile(file_path, "r") as zf:
                names = zf.namelist()
                for req in ["[Content_Types].xml", "word/document.xml"]:
                    if req not in names:
                        errors.append(f"Missing: {req}")
                for name in names:
                    if name.endswith((".xml", ".rels")):
                        try:
                            lxml.etree.fromstring(zf.read(name))
                        except lxml.etree.XMLSyntaxError as e:
                            errors.append(f"{name}: {e}")
        except zipfile.BadZipFile:
            errors.append("Not a valid ZIP archive")
        return errors
